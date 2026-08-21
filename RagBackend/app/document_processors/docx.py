"""
DOCX extraction with heading hierarchy and inline table support.
"""

from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import Docx2txtLoader

from app.config import get_logger
from app.document_processors.base import format_table_as_markdown

logger = get_logger("askdoc.processors.docx")


def extract_docx_with_structure(docx_path: str, filename: str) -> List[Document]:
    """
    Walks a DOCX document in sequential element order (paragraphs + tables):
    1. Converts Heading 1/2/3 styles to Markdown # / ## / ### headers inline.
    2. Converts tables to inlined Markdown tables in their exact document position.
    3. Tracks a running heading_path stack (e.g. 'Introduction > Background')
       and attaches it to section Documents before chunking.
    """
    try:
        import docx
        from docx.text.paragraph import Paragraph as DocxParagraph
        from docx.table import Table as DocxTable
    except ImportError:
        logger.warning("python-docx not available, falling back to Docx2txtLoader")
        loader = Docx2txtLoader(docx_path)
        return loader.load()

    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        logger.warning(f"Error opening docx with python-docx ({e}), falling back to Docx2txtLoader")
        loader = Docx2txtLoader(docx_path)
        return loader.load()

    documents = []
    current_section_lines = []
    heading_stack = []  # [(level, title)] e.g. [(1, "Introduction"), (2, "Background")]

    def get_heading_path():
        return " > ".join([title for _, title in heading_stack]) if heading_stack else ""

    def flush_section():
        if current_section_lines:
            content = "\n".join(current_section_lines).strip()
            if content:
                path = get_heading_path()
                documents.append(Document(
                    page_content=content,
                    metadata={
                        "source": filename,
                        "title": filename,
                        "heading_path": path,
                        "section": path
                    }
                ))
            current_section_lines.clear()

    # Iterate over all child elements of the document body in order
    for child in doc.element.body:
        if child.tag.endswith('p'):
            p = DocxParagraph(child, doc)
            text = p.text.strip()
            if not text:
                continue

            style_name = (p.style.name if p.style and p.style.name else "").lower()
            
            # Detect heading levels
            if "heading 1" in style_name or style_name == "title":
                flush_section()
                heading_stack = [(1, text)]
                current_section_lines.append(f"# {text}\n")
            elif "heading 2" in style_name or style_name == "subtitle":
                flush_section()
                heading_stack = [(lvl, t) for lvl, t in heading_stack if lvl < 2]
                heading_stack.append((2, text))
                current_section_lines.append(f"## {text}\n")
            elif "heading 3" in style_name:
                flush_section()
                heading_stack = [(lvl, t) for lvl, t in heading_stack if lvl < 3]
                heading_stack.append((3, text))
                current_section_lines.append(f"### {text}\n")
            elif "heading 4" in style_name:
                flush_section()
                heading_stack = [(lvl, t) for lvl, t in heading_stack if lvl < 4]
                heading_stack.append((4, text))
                current_section_lines.append(f"#### {text}\n")
            else:
                current_section_lines.append(text)

        elif child.tag.endswith('tbl'):
            tbl = DocxTable(child, doc)
            table_data = []
            for row in tbl.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    table_data.append(row_cells)
                    
            if table_data:
                md_table = format_table_as_markdown(table_data)
                if md_table.strip():
                    current_section_lines.append(md_table)

    flush_section()

    if not documents:
        # Fallback if whole doc had no structured headings
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if full_text.strip():
            documents.append(Document(
                page_content=full_text.strip(),
                metadata={"source": filename, "title": filename, "heading_path": "", "section": ""}
            ))

    return documents
